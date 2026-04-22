# Copyright (c) 2026 Marco De Roni. All rights reserved.
# Licensed under the MIT License — see LICENSE file for details.

import os
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(
    contract_name: str,
    analysis: str,
    provider: str,
    output_dir: str,
    pii_summary: dict = None
) -> str:

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading("CONTRACT REVIEW REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"Document: {contract_name}\n"
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}\n"
        f"Powered by: {provider.upper()}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- PII Redaction Summary ---
    if pii_summary and pii_summary.get("total_entities", 0) > 0:
        doc.add_heading("Privacy & PII Redaction Summary", level=1)
        total = pii_summary.get("total_entities", 0)
        breakdown = pii_summary.get("breakdown", {})

        doc.add_paragraph(
            f"Before analysis, {total} sensitive entities were automatically "
            f"redacted and restored in this report."
        )

        if breakdown:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Entity Type"
            table.rows[0].cells[1].text = "Count"
            for entity_type, count in sorted(breakdown.items()):
                row = table.add_row().cells
                row[0].text = entity_type
                row[1].text = str(count)

        doc.add_paragraph()

    # --- Analysis content ---
    for line in analysis.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=2)
        elif line.startswith("- 🔴"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.replace("- ", ""))
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.size = Pt(10)
        elif line.startswith("- 🟡"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.replace("- ", ""))
            run.font.color.rgb = RGBColor(0xFF, 0xC0, 0x00)
            run.font.size = Pt(10)
        elif line.startswith("- 🟢"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.replace("- ", ""))
            run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            run.font.size = Pt(10)
        elif line.startswith("- ⚪"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.replace("- ", ""))
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.font.size = Pt(10)
        elif line.startswith("- "):
            doc.add_paragraph(line, style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    os.makedirs(output_dir, exist_ok=True)
    safe_name = contract_name.replace(" ", "_").replace("/", "_")
    path = os.path.join(output_dir, f"report_{safe_name}.docx")
    doc.save(path)
    return path


def generate_pdf(docx_path: str) -> str:
    """Converts Word report to PDF."""
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace(".docx", ".pdf")
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        print(f"   ⚠️  PDF conversion failed: {e}")
        return None


def generate_report(
    contract_name: str,
    analysis: str,
    provider: str,
    output_dir: str,
    output_format: str = "both",
    pii_summary: dict = None
) -> dict:
    """Generates Word and/or PDF report."""
    paths = {}

    docx_path = generate_docx(contract_name, analysis, provider, output_dir, pii_summary=pii_summary)
    paths["docx"] = docx_path

    if output_format in ("pdf", "both"):
        pdf_path = generate_pdf(docx_path)
        if pdf_path:
            paths["pdf"] = pdf_path

    return paths